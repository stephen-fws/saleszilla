"""Extract text from chat file attachments.

Supports:
  - PDF (via pypdf)
  - DOCX (via python-docx)
  - Plain text family: txt, md, csv, json, log, py, js, ts, tsx, jsx, html, sql,
    yaml, yml, xml, sh, ini, cfg, toml

Each extracted file is wrapped in a structured marker so the chat agent can
clearly attribute content to a source file.
"""

from __future__ import annotations

import io
import logging
from typing import Iterable

logger = logging.getLogger(__name__)


# File size guard — protect against accidentally pasting a 500 MB file
MAX_FILE_BYTES = 25 * 1024 * 1024  # 25 MB
# Truncate extracted text per file (Claude has a context budget)
MAX_TEXT_PER_FILE = 200_000  # ~50K tokens

TEXT_EXTENSIONS = {
    "txt", "md", "csv", "json", "log", "py", "js", "ts", "tsx", "jsx", "html",
    "htm", "css", "sql", "yaml", "yml", "xml", "sh", "ini", "cfg", "toml",
    "rtf", "tsv",
}


def _ext(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _truncate(text: str) -> str:
    if len(text) <= MAX_TEXT_PER_FILE:
        return text
    return text[:MAX_TEXT_PER_FILE] + f"\n\n[... truncated, original was {len(text):,} chars]"


def _extract_pdf(content: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[PDF support not available — please install pypdf]"
    try:
        reader = PdfReader(io.BytesIO(content))
        chunks: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            if text.strip():
                chunks.append(f"--- Page {i + 1} ---\n{text.strip()}")
        return "\n\n".join(chunks) if chunks else "[PDF contained no extractable text]"
    except Exception as e:
        logger.warning("PDF extraction failed: %s", e)
        return f"[PDF extraction failed: {e}]"


def _extract_docx(content: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return "[DOCX support not available — please install python-docx]"
    try:
        document = docx.Document(io.BytesIO(content))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        # Also extract tables
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        return "\n".join(paragraphs) if paragraphs else "[DOCX contained no text]"
    except Exception as e:
        logger.warning("DOCX extraction failed: %s", e)
        return f"[DOCX extraction failed: {e}]"


def _extract_text_file(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return "[Could not decode text file]"


def extract_attachment(filename: str, content: bytes) -> tuple[bool, str]:
    """Extract text from a single file. Returns (success, text_or_error)."""
    if len(content) > MAX_FILE_BYTES:
        return False, f"[File '{filename}' too large ({len(content):,} bytes, limit {MAX_FILE_BYTES:,})]"

    ext = _ext(filename)
    if ext == "pdf":
        text = _extract_pdf(content)
    elif ext in ("docx", "doc"):
        if ext == "doc":
            return False, f"[Legacy .doc files not supported — please save '{filename}' as .docx]"
        text = _extract_docx(content)
    elif ext in TEXT_EXTENSIONS:
        text = _extract_text_file(content)
    else:
        return False, f"[Unsupported file type: '{filename}'. Supported: PDF, DOCX, TXT, CSV, MD, JSON, code files]"

    return True, _truncate(text)


def build_attachment_block(filename: str, text: str) -> str:
    """Format an extracted file as a labelled block to embed in the user message."""
    return (
        f"\n\n=== ATTACHED FILE: {filename} ===\n"
        f"{text}\n"
        f"=== END OF FILE: {filename} ===\n"
    )


def enrich_message_with_attachments(
    user_message: str,
    files: Iterable[tuple[str, bytes]],
) -> tuple[str, list[dict]]:
    """Take the raw user message and a list of (filename, bytes) tuples.

    Returns:
      - enriched_message: the user's text + extracted file content blocks appended
      - file_summaries: list of dicts describing each file's processing result
    """
    summaries: list[dict] = []
    extracted_blocks: list[str] = []

    for filename, content in files:
        ok, text = extract_attachment(filename, content)
        summaries.append({
            "filename": filename,
            "size": len(content),
            "ok": ok,
            "chars": len(text) if ok else 0,
            "error": None if ok else text,
        })
        if ok:
            extracted_blocks.append(build_attachment_block(filename, text))
        else:
            # Still tell Claude that an upload failed
            extracted_blocks.append(f"\n\n[Note: failed to read uploaded file '{filename}': {text}]\n")

    enriched = user_message + "".join(extracted_blocks)
    return enriched, summaries
